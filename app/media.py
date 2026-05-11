"""Process uploaded files into a normalized form for LLM consumption.

- Images  -> base64 + mime type (kept for vision-capable providers)
- Docs    -> extracted plain text (PDF, DOCX, XLSX, HTML, MD, TXT, CSV, JSON)
- Unknown -> attempt utf-8 decode, otherwise note as binary
"""
from __future__ import annotations

import base64
import io
import json
from dataclasses import dataclass, field
from typing import List, Optional

from fastapi import UploadFile

from .schemas import AttachmentInfo
from .utils.logger import logger

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None

try:
    import docx  # python-docx
except ImportError:  # pragma: no cover
    docx = None

try:
    from openpyxl import load_workbook
except ImportError:  # pragma: no cover
    load_workbook = None

try:
    from bs4 import BeautifulSoup
except ImportError:  # pragma: no cover
    BeautifulSoup = None

try:
    from PIL import Image
except ImportError:  # pragma: no cover
    Image = None


IMAGE_MIME_TYPES = {
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/webp",
    "image/gif",
    "image/heic",
    "image/heif",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".heic", ".heif", ".bmp"}

# Reasonable limits to keep request bodies manageable
MAX_DOC_TEXT_CHARS = 60_000
MAX_IMAGE_BYTES = 10 * 1024 * 1024  # 10 MB


@dataclass
class ProcessedImage:
    filename: str
    mime_type: str
    data_base64: str
    size_bytes: int


@dataclass
class ProcessedDocument:
    filename: str
    mime_type: str
    text: str
    size_bytes: int


@dataclass
class ProcessedMedia:
    images: List[ProcessedImage] = field(default_factory=list)
    documents: List[ProcessedDocument] = field(default_factory=list)
    info: List[AttachmentInfo] = field(default_factory=list)

    @property
    def has_images(self) -> bool:
        return len(self.images) > 0

    @property
    def combined_document_text(self) -> str:
        if not self.documents:
            return ""
        chunks: List[str] = []
        for doc in self.documents:
            chunks.append(f"--- File: {doc.filename} ({doc.mime_type}) ---\n{doc.text}")
        return "\n\n".join(chunks)


def _ext(filename: str) -> str:
    if "." not in filename:
        return ""
    return filename[filename.rfind(".") :].lower()


def _is_image(filename: str, content_type: str) -> bool:
    if content_type and content_type.lower() in IMAGE_MIME_TYPES:
        return True
    return _ext(filename) in IMAGE_EXTENSIONS


def _truncate(text: str, limit: int = MAX_DOC_TEXT_CHARS) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + f"\n\n[... truncated {len(text) - limit} chars ...]"


def _extract_pdf(data: bytes) -> str:
    if PdfReader is None:
        return "[pypdf not installed - cannot read PDF]"
    try:
        reader = PdfReader(io.BytesIO(data))
        parts: List[str] = []
        for i, page in enumerate(reader.pages):
            try:
                parts.append(page.extract_text() or "")
            except Exception as e:
                parts.append(f"[page {i + 1} extraction failed: {e}]")
        return _truncate("\n".join(parts).strip())
    except Exception as e:
        logger.warning(f"PDF parse failed: {e}")
        return f"[PDF parse error: {e}]"


def _extract_docx(data: bytes) -> str:
    if docx is None:
        return "[python-docx not installed - cannot read DOCX]"
    try:
        document = docx.Document(io.BytesIO(data))
        text = "\n".join(p.text for p in document.paragraphs if p.text)
        for table in document.tables:
            for row in table.rows:
                text += "\n" + " | ".join(cell.text for cell in row.cells)
        return _truncate(text.strip())
    except Exception as e:
        logger.warning(f"DOCX parse failed: {e}")
        return f"[DOCX parse error: {e}]"


def _extract_xlsx(data: bytes) -> str:
    if load_workbook is None:
        return "[openpyxl not installed - cannot read XLSX]"
    try:
        wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
        parts: List[str] = []
        for ws in wb.worksheets:
            parts.append(f"# Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = ["" if c is None else str(c) for c in row]
                if any(cells):
                    parts.append(" | ".join(cells))
        return _truncate("\n".join(parts).strip())
    except Exception as e:
        logger.warning(f"XLSX parse failed: {e}")
        return f"[XLSX parse error: {e}]"


def _extract_html(data: bytes) -> str:
    if BeautifulSoup is None:
        return data.decode("utf-8", errors="replace")
    try:
        soup = BeautifulSoup(data, "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        return _truncate(soup.get_text(separator="\n").strip())
    except Exception as e:
        logger.warning(f"HTML parse failed: {e}")
        return data.decode("utf-8", errors="replace")


def _extract_text_like(data: bytes) -> str:
    try:
        return _truncate(data.decode("utf-8"))
    except UnicodeDecodeError:
        try:
            return _truncate(data.decode("latin-1"))
        except Exception:
            return "[binary content - could not decode]"


def _extract_json(data: bytes) -> str:
    try:
        obj = json.loads(data.decode("utf-8"))
        return _truncate(json.dumps(obj, indent=2, ensure_ascii=False))
    except Exception:
        return _extract_text_like(data)


def _process_document(filename: str, content_type: str, data: bytes) -> Optional[str]:
    ext = _ext(filename)
    ct = (content_type or "").lower()

    if ext == ".pdf" or "pdf" in ct:
        return _extract_pdf(data)
    if ext == ".docx" or "officedocument.wordprocessingml" in ct:
        return _extract_docx(data)
    if ext in {".xlsx", ".xlsm"} or "spreadsheetml" in ct:
        return _extract_xlsx(data)
    if ext in {".html", ".htm"} or "html" in ct:
        return _extract_html(data)
    if ext == ".json" or "json" in ct:
        return _extract_json(data)
    if ext in {".txt", ".md", ".markdown", ".csv", ".tsv", ".log", ".yaml", ".yml", ".xml", ".ini", ".cfg"}:
        return _extract_text_like(data)
    if ct.startswith("text/"):
        return _extract_text_like(data)
    return None


async def process_uploads(files: Optional[List[UploadFile]]) -> ProcessedMedia:
    out = ProcessedMedia()
    if not files:
        return out

    for f in files:
        if f is None or not getattr(f, "filename", None):
            continue
        data = await f.read()
        size = len(data)
        ct = (f.content_type or "").lower()
        filename = f.filename or "upload"

        if _is_image(filename, ct):
            if size > MAX_IMAGE_BYTES:
                logger.warning(f"Image {filename} exceeds {MAX_IMAGE_BYTES} bytes, skipping")
                out.info.append(
                    AttachmentInfo(
                        filename=filename,
                        content_type=ct or "image/*",
                        size_bytes=size,
                        kind="image",
                    )
                )
                continue
            mime = ct or "image/png"
            if not mime.startswith("image/"):
                mime = "image/png"
            b64 = base64.b64encode(data).decode("ascii")
            out.images.append(
                ProcessedImage(
                    filename=filename,
                    mime_type=mime,
                    data_base64=b64,
                    size_bytes=size,
                )
            )
            out.info.append(
                AttachmentInfo(
                    filename=filename,
                    content_type=mime,
                    size_bytes=size,
                    kind="image",
                )
            )
            continue

        text = _process_document(filename, ct, data)
        if text is not None:
            out.documents.append(
                ProcessedDocument(
                    filename=filename,
                    mime_type=ct or "application/octet-stream",
                    text=text,
                    size_bytes=size,
                )
            )
            out.info.append(
                AttachmentInfo(
                    filename=filename,
                    content_type=ct or "application/octet-stream",
                    size_bytes=size,
                    kind="document" if not ct.startswith("text/") else "text",
                    extracted_chars=len(text),
                )
            )
            continue

        decoded = _extract_text_like(data)
        if decoded and not decoded.startswith("[binary"):
            out.documents.append(
                ProcessedDocument(
                    filename=filename,
                    mime_type=ct or "text/plain",
                    text=decoded,
                    size_bytes=size,
                )
            )
            out.info.append(
                AttachmentInfo(
                    filename=filename,
                    content_type=ct or "text/plain",
                    size_bytes=size,
                    kind="text",
                    extracted_chars=len(decoded),
                )
            )
        else:
            out.info.append(
                AttachmentInfo(
                    filename=filename,
                    content_type=ct or "application/octet-stream",
                    size_bytes=size,
                    kind="unknown",
                )
            )
    return out
